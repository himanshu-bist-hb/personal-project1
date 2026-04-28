"""
Generate a polished Word document of the BA Rate Pages developer guide,
suitable for sharing with stakeholders.

Run:  python generate_guide_docx.py
Output: BA_Rate_Pages_Developer_Guide.docx
"""
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

ACCENT      = RGBColor(0x1F, 0x3A, 0x68)   # deep navy
SUBACCENT   = RGBColor(0x2E, 0x5C, 0x9E)
TEXT_DARK   = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_MUTED  = RGBColor(0x55, 0x55, 0x55)
CODE_BG     = "F4F6F8"
HEAD_BG     = "1F3A68"
ZEBRA_BG    = "F7F9FC"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def set_paragraph_border(paragraph, color="DCE0E5", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), color)
        pBdr.append(b)
    p_pr.append(pBdr)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A68")
    pBdr.append(bottom)
    p_pr.append(pBdr)


def add_code_block(doc, code_text):
    """Add a monospace, shaded, bordered block — looks like a styled code listing."""
    for line in code_text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_DARK
        # Force monospace on East-Asian as well
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:cs"), "Consolas")
        rPr.append(rFonts)
        set_paragraph_shading(p, CODE_BG)
        set_paragraph_border(p)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_inline_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB8, 0x2A, 0x2A)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = SUBACCENT
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TEXT_DARK
    return p


def add_para(doc, text, *, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_callout(doc, label, text, color_hex="FFF4E0", border_hex="E8B45C"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    label_run = p.add_run(f"{label}  ")
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor(0x8A, 0x57, 0x10)
    body_run = p.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = TEXT_DARK
    set_paragraph_shading(p, color_hex)
    set_paragraph_border(p, color=border_hex, size="8")


def add_table(doc, headers, rows, col_widths_inches=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    if col_widths_inches:
        for i, w in enumerate(col_widths_inches):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        set_cell_shading(hdr[i], HEAD_BG)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            for chunk in _parse_inline(str(val)):
                kind, txt = chunk
                run = para.add_run(txt)
                run.font.name = "Consolas" if kind == "code" else "Calibri"
                run.font.size = Pt(10)
                if kind == "code":
                    run.font.color.rgb = RGBColor(0xB8, 0x2A, 0x2A)
                if kind == "bold":
                    run.font.bold = True
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r_idx % 2 == 1:
                set_cell_shading(cells[i], ZEBRA_BG)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def _parse_inline(text):
    """Tiny inline parser: `code` and **bold** segments."""
    out = []
    i = 0
    while i < len(text):
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j == -1:
                out.append(("plain", text[i:])); break
            out.append(("code", text[i + 1:j])); i = j + 1
        elif text[i:i+2] == "**":
            j = text.find("**", i + 2)
            if j == -1:
                out.append(("plain", text[i:])); break
            out.append(("bold", text[i + 2:j])); i = j + 2
        else:
            j = i
            while j < len(text) and text[j] != "`" and text[j:j+2] != "**":
                j += 1
            out.append(("plain", text[i:j])); i = j
    return out


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page setup: letter, modest margins
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ====================================================================
    # COVER PAGE
    # ====================================================================
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(120)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("BA Rate Pages")
    r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Developer Guide")
    r.font.size = Pt(22); r.font.color.rgb = SUBACCENT

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_before = Pt(20)
    r = tag.add_run("Adding rules · Headings & subtitles · Page-break rules · PDF generation")
    r.font.size = Pt(12); r.font.italic = True; r.font.color.rgb = TEXT_MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(180)
    r = meta.add_run(f"Nationwide Insurance  ·  BA Analytics Division\nInternal Use Only\nPrepared: {date.today().strftime('%B %d, %Y')}")
    r.font.size = Pt(10); r.font.color.rgb = TEXT_MUTED

    doc.add_page_break()

    # ====================================================================
    # TABLE OF CONTENTS
    # ====================================================================
    add_h1(doc, "Table of Contents")
    add_horizontal_rule(doc)
    toc_rows = [
        ("1.", "The 30-second tour"),
        ("2.", "The files you will touch"),
        ("3.", "How a rate page is born — full lifecycle"),
        ("4.", "PART 1 — Adding a new rule (new Excel tab)"),
        ("5.", "PART 2 — Headings, subtitles, sub-headings"),
        ("6.", "PART 3 — Page-break rules"),
        ("7.", "PART 4 — PDF generation"),
        ("8.", "Common mistakes (read this!)"),
        ("9.", "Cheat sheet — where to find what"),
    ]
    for num, label in toc_rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{num}  ")
        r1.font.bold = True; r1.font.color.rgb = ACCENT; r1.font.size = Pt(12)
        r2 = p.add_run(label)
        r2.font.size = Pt(12); r2.font.color.rgb = TEXT_DARK

    add_para(doc, "")
    add_para(doc, "If you have never touched the codebase before, start at Section 1. Every step has a copy-paste-able example.", italic=True, color=TEXT_MUTED, size=10)

    doc.add_page_break()

    # ====================================================================
    # SECTION 1 — 30-SECOND TOUR
    # ====================================================================
    add_h1(doc, "1.  The 30-second tour")
    add_horizontal_rule(doc)
    add_para(doc, "The application reads Rate Books (Excel files containing insurance rate tables) and produces a single Rate Pages workbook (Excel + PDF) that is filed with the Department of Insurance.")
    add_para(doc, "The pipeline:")
    add_code_block(doc, """\
Streamlit UI (app.py)
        │  user uploads ratebooks, picks save folder, clicks Run
        ▼
BARatePages.run(...)            ← orchestration: opens files, calls everything
        │
        ▼
BARates.Auto.buildBAPages()     ← the BIG function — for each rule:
        │                          1. compares companies (clustering)
        │                          2. builds a DataFrame
        │                          3. tells the Excel factory to make a sheet
        ▼
ExcelSettingsBA.Excel           ← generateWorksheet*() creates the tab
        │                          formatWorksheet*() styles it
        ▼
   .xlsx is saved
        │
        ▼
BApagebreaks.process_pagebreaks ← applies page breaks per rule
        │
        ▼
   user clicks "Generate PDF Document"
        ▼
BApagebreaks.export_to_pdf      ← drives Excel via COM to make a real PDF
        │
        ▼
   .pdf is saved next to the .xlsx""")
    add_para(doc, "You will rarely touch app.py or BARatePages.py. The interesting work happens in BARates.py, ExcelSettingsBA.py, and BApagebreaks.py.")

    # ====================================================================
    # SECTION 2 — FILES YOU WILL TOUCH
    # ====================================================================
    add_h1(doc, "2.  The files you will touch")
    add_horizontal_rule(doc)
    add_table(doc,
        ["File", "What it does", "When to touch it"],
        [
            ["`BARates.py`",
             "The brain. Holds all rule logic, builds DataFrames, decides what goes on each sheet. Look inside `buildBAPages()`.",
             "Adding a new rule, changing what data shows on a sheet, changing the page title."],
            ["`ExcelSettingsBA.py`",
             "The Excel factory. Knows how to draw a tab — title in A1, subtitle in A2, table starts at A4, borders, fonts, headers, footers.",
             "Only when you need a layout shape that does not already exist. 99% of the time you reuse the existing methods."],
            ["`BApagebreaks.py`",
             "Post-processing. Decides where each printed page should break, what the print area is, portrait vs landscape, etc.",
             "Whenever you add a new rule and want it to print the way your team expects."],
            ["`config/constants.py`",
             "Single source of truth: margins, fonts, paths, company names.",
             "Only when changing things that affect the whole workbook (e.g. font size 10 → 11)."],
        ],
        col_widths_inches=[1.6, 3.0, 2.4])
    add_para(doc, "You will not touch:")
    add_bullet(doc, "BARatePages.py — orchestration, already wired up.")
    add_bullet(doc, "app.py — UI, already wired up.")
    add_bullet(doc, "BA Input File.xlsx — read-only dependency.")

    # ====================================================================
    # SECTION 3 — LIFECYCLE
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "3.  How a rate page is born — full lifecycle")
    add_horizontal_rule(doc)
    add_para(doc, "For one rule (e.g. Rule 208 Expense Constant), here is exactly what happens, step by step:")
    add_code_block(doc, """\
1. buildBAPages() reaches the Rule 208 block
2. self.compareCompanies('ExpenseConstant_Ext')
        ─► figures out which companies share the same numbers
3. for each cluster of companies:
       a) self.buildExpenseConstant(comp_name)
              ─► returns a pandas DataFrame
       b) RatePages.generateWorksheet(
              wsTitle       = 'Rule 208',                      # Excel TAB name
              tableTitle    = 'RULE 208. EXPENSE CONSTANT',    # cell A1
              tableSubtitle = '208.B. Rate and Premium ...',   # cell A2
              df            = the DataFrame from step (a),
              useIndex      = False,
              useHeader     = True)
       c) self.overideFooter(ws, CompanyTest)
              ─► writes the right company names in the footer
4. After all rules done: RatePages.createIndex()
        ─► builds the Index tab with one hyperlink per rule
5. The workbook is saved as .xlsx
6. process_pagebreaks(xlsx_out, pdf_out)
        ─► runs through every sheet, applies page-break rules
7. (Optional) user clicks "Generate PDF Document"
        ─► export_to_pdf(xlsx, pdf) drives Excel and exports""")
    add_para(doc, "The two things that need your attention when adding a rule are:")
    add_bullet(doc, "Step 3b — picking the right generateWorksheet* method.")
    add_bullet(doc, "Step 6 — registering a page-break rule for the new tab.")

    # ====================================================================
    # SECTION 4 — PART 1: ADDING A NEW RULE
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "4.  PART 1 — Adding a new rule (new Excel tab)")
    add_horizontal_rule(doc)

    add_h2(doc, "4.1  The pattern")
    add_para(doc, "Every rule block in buildBAPages() (in BARates.py) follows the same six-line pattern. Only the bold parts change between rules:")
    add_code_block(doc, """\
# Rule XXX
self.compareCompanies('SomeRateTableName_Ext')        # cluster companies
for CompanyTest in self.CompanyListDif:
    comp_name = self.extract_company_name(CompanyTest)
    self.title_company_name = CompanyTest
    if len(self.CompanyListDif) == 1:
        self.title_company_name = ""                  # only one cluster → no suffix

    RatePages.generateWorksheet(
        'Rule XXX ' + self.title_company_name,                       # tab name
        'RULE XXX. SHORT DESCRIPTION ' + self.title_company_name,    # cell A1
        'XXX.A. Rate Computation',                                   # cell A2
        self.buildRuleXXX(comp_name),                                # the DataFrame
        False,    # useIndex
        True      # useHeader
    )
    self.overideFooter(RatePages.getWB()['Rule XXX ' + self.title_company_name], CompanyTest)""")

    add_h2(doc, "4.2  Step-by-step: adding 'Rule 999 — Imaginary Coverage'")

    add_h3(doc, "Step 1 — Build the DataFrame")
    add_para(doc, "Add a method to the Auto class (anywhere in BARates.py) that returns a pandas DataFrame shaped like the table you want to print:")
    add_code_block(doc, """\
def buildRule999(self, comp_name):
    # Use self.buildDataFrame('SomeTableName_Ext') to fetch a nested rate table,
    # or build the DataFrame from scratch if it is a static table.
    df = self.buildDataFrame('ImaginaryFactor_Ext')
    df.columns = ['Coverage', 'Factor']
    return df""")
    add_callout(doc, "TIP", "Look at buildExpenseConstant, build231C, or buildAntiqueAutoLiabFactors for short, readable templates.")

    add_h3(doc, "Step 2 — Wire it into buildBAPages()")
    add_code_block(doc, """\
# Rule 999
self.compareCompanies('ImaginaryFactor_Ext')
for CompanyTest in self.CompanyListDif:
    comp_name = self.extract_company_name(CompanyTest)
    self.title_company_name = CompanyTest
    if len(self.CompanyListDif) == 1:
        self.title_company_name = ""

    RatePages.generateWorksheet(
        'Rule 999 ' + self.title_company_name,
        'RULE 999. IMAGINARY COVERAGE ' + self.title_company_name,
        '999.A. Premium Computation',
        self.buildRule999(comp_name),
        False, True
    )
    self.overideFooter(RatePages.getWB()['Rule 999 ' + self.title_company_name], CompanyTest)""")

    add_h3(doc, "Step 3 — Add a format pass at the bottom of buildBAPages()")
    add_para(doc, "After all generateWorksheet* calls, there is a long block where each rule gets a styling pass. Add:")
    add_code_block(doc, """\
Rule999_Sheets = [name for name in excel_Sheet_names if name.startswith('Rule 999')]
for Rule in Rule999_Sheets:
    self.formatWorksheet(AutoPages[Rule])""")

    add_h3(doc, "Step 4 — Add a page-break rule")
    add_para(doc, "See Section 6 (Page-break rules).")

    add_h2(doc, "4.3  Picking the right generateWorksheet* method")
    add_para(doc, "Pick the layout that matches what you are drawing:")
    add_table(doc,
        ["If your sheet has…", "Use this method"],
        [
            ["1 table", "`generateWorksheet(title, A1, A2, df, useIndex, useHeader)`"],
            ["2 tables, each with its own subtitle", "`generateWorksheet2tables(title, A1, A2, df1, A2_for_df2, df2, ...)`"],
            ["2–14 tables, one shared subtitle, stacked", "`generateWorksheet2tbls`, `generateWorksheet3tables`, … `generateWorksheet14tables`"],
            ["Many tables, each with its own subtitle, dynamic count", "`generateWorksheetTablesX(title, A1, [subtitles], [dfs], ...)`"],
            ["2 side-by-side tables (Rule 23B style)", "`generateWorksheet23B(...)`"],
            ["Rule 222 fixed-row layout", "`generateRule222(...)`"],
            ["Add a 2nd table to an existing tab (Fleet style)", "`ModifyFleetTable(sheetName, subtitle, df, ...)`"],
        ],
        col_widths_inches=[3.0, 4.0])
    add_callout(doc, "UNDER THE HOOD",
                "generateWorksheet2tbls through generateWorksheet14tables all delegate to one helper — generate_stacked_tables(...). They exist as named aliases so older code reads naturally. You can also call generate_stacked_tables directly with a list of dataframes.",
                color_hex="EAF1FB", border_hex="9FB7DA")

    add_h2(doc, "4.4  The two booleans: useIndex and useHeader")
    add_bullet(doc, "useIndex = True  →  write the DataFrame's row index as the first column. Almost always False for rate pages.")
    add_bullet(doc, "useHeader = True  →  write the DataFrame's column names as a header row. Almost always True.")
    add_para(doc, "If unsure: False, True.", italic=True, color=TEXT_MUTED)

    add_h2(doc, "4.5  Do not forget overideFooter")
    add_para(doc, "Every generateWorksheet* call writes a default footer. But when companies are clustered (e.g. NGIC + NACO share a tab), the footer must list the right names for that cluster. overideFooter does that:")
    add_code_block(doc, """\
self.overideFooter(RatePages.getWB()['Rule 999 ' + self.title_company_name], CompanyTest)""")
    add_callout(doc, "WARNING",
                "Skipping overideFooter leaves the default footer (which assumes all companies). The DoI filing will list the wrong names. Always include it.",
                color_hex="FFE8E8", border_hex="D88A8A")

    # ====================================================================
    # SECTION 5 — PART 2: HEADINGS & SUBHEADINGS
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "5.  PART 2 — Headings, subtitles, sub-headings")
    add_horizontal_rule(doc)
    add_para(doc, "The Excel layout for every rate page tab is:")
    add_code_block(doc, """\
+----------------------------------------------------------+
| Row 1: tableTitle           ← A1, BOLD, big              |
| Row 2: tableSubtitle        ← A2, italic                 |
| Row 3: (blank)                                           |
| Row 4: column headers ┐                                  |
| Row 5+: data rows     ├── the DataFrame                  |
|                       ┘                                  |
+----------------------------------------------------------+
| Header (top of every printed page):                      |
|   Left:   "Commercial Lines Manual: Division One - Auto" |
|   Center: "{State} - Rate Exceptions"                    |
|   Right:  "Effective: New {date}  Renewal {date}"        |
| Footer (bottom of every page): company names + page #    |
+----------------------------------------------------------+""")

    add_h2(doc, "5.1  The main heading (cell A1)")
    add_para(doc, "Third argument to generateWorksheet:")
    add_code_block(doc, """\
RatePages.generateWorksheet(
    'Rule 208',                              # tab name
    'RULE 208. EXPENSE CONSTANT',            # ← THIS goes in cell A1 (bold)
    '208.B. Rate and Premium Computation',   # subtitle
    df, False, True
)""")
    add_para(doc, "Convention used across the project: \"RULE NNN. SHORT DESCRIPTION\". To include the company suffix on the heading, append self.title_company_name.")

    add_h2(doc, "5.2  The subtitle (cell A2)")
    add_para(doc, "Fourth argument. To skip the subtitle, pass exactly ' ' (a single space). The factory has special-case logic: when tableSubtitle == ' ', the data starts on row 3 instead of row 4 — useful for compact base-rate pages.")
    add_code_block(doc, """\
RatePages.generateWorksheet(
    'Rule 222 TTT BR', 'RULE 222. … BASE RATES',
    ' ',                                # ← no subtitle
    df, False, True
)""")

    add_h2(doc, "5.3  Multiple sub-headings (one per table)")
    add_h3(doc, "A) Two tables with named subtitles")
    add_code_block(doc, """\
RatePages.generateWorksheet2tables(
    'Rule 232 B',
    'RULE 232. PREMIUM DEVELOPMENT - PRIVATE PASSENGER TYPES',  # A1
    '232.B.1.b. Liability Fleet Size Factors',                  # subtitle for df1
    self.buildPPTLiabFleetFactors(comp_name),                   # df1
    '232.B.4.d. Physical Damage Fleet Size Factors',            # subtitle for df2
    self.buildPPTPhysDamFleetFactors(comp_name),                # df2
    False, True
)""")

    add_h3(doc, "B) Variable number of sub-tables")
    add_para(doc, "Use generateWorksheetTablesX. Subtitles and dataframes are passed as parallel lists:")
    add_code_block(doc, """\
subtitles = [
    '275.B.1.(a).(2). Trucks, Tractors, and Trailers Factors',
    '275.B.1.(a).(2). Private Passenger Types Factor',
    '275.B.1.(a).(2). Motorcycles Factors',
]
tables = [
    self.buildSomething1(comp_name),
    self.buildSomething2(comp_name),
    self.buildSomething3(comp_name),
]
RatePages.generateWorksheetTablesX(
    'Rule 275',
    'RULE 275. LEASING OR RENTAL CONCERNS',     # A1
    subtitles, tables,
    False, True
)""")

    add_h2(doc, "5.4  Plain stacked tables (one shared subtitle)")
    add_code_block(doc, """\
RatePages.generateWorksheet3tables(
    'Rule 222 B',
    'RULE 222. PREMIUM DEVELOPMENT - TRUCK, TRACTOR, TRAILER',
    '222.B.1.a. Liability Fleet Size Factors',          # one subtitle for the whole sheet
    df1, df2, df3,
    False, True
)""")
    add_para(doc, "Methods exist for 2 through 14 stacked dataframes; they only differ in the number of df arguments accepted.")

    add_h2(doc, "5.5  Header & footer")
    add_para(doc, "These are applied automatically. To change:")
    add_table(doc,
        ["What you want to change", "Where to change it"],
        [
            ["Center / right header text", "`_apply_standard_header` in ExcelSettingsBA.py"],
            ["Left header text (global)", "`HEADER_LEFT_TEXT` in config/constants.py"],
            ["Footer company names per cluster", "`overideFooter` in BARates.py"],
            ["Page number / state / tab name in center footer", "Already wired: `{StateAbb} - &[Tab] - &P`"],
        ],
        col_widths_inches=[3.2, 3.8])

    # ====================================================================
    # SECTION 6 — PART 3: PAGE-BREAK RULES
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "6.  PART 3 — Page-break rules")
    add_horizontal_rule(doc)
    add_para(doc, "All page-break logic lives in BApagebreaks.py. You only ever edit that one file for page-break work.")

    add_h2(doc, "6.1  The mental model")
    add_para(doc, "After the workbook is saved, process_pagebreaks(xlsx_path, pdf_path) opens it with openpyxl (no Excel app needed) and walks every sheet. For each sheet it:")
    add_bullet(doc, "Sets defaults: print_title_rows = \"1:1\" (row 1 repeats on every printed page) and fit_single_page (whole sheet on one page).")
    add_bullet(doc, "Walks the SHEET_RULES list and runs the first handler whose prefix matches the sheet name.")
    add_para(doc, "A page-break rule is just a Python function that decides how a sheet should print. Two steps to add one: write the handler, register it in SHEET_RULES.")

    add_h2(doc, "6.2  The four helpers you will use 95% of the time")
    add_table(doc,
        ["Helper", "What it does", "When to use"],
        [
            ["`fit_single_page(ws)`", "Shrinks the whole sheet onto one printed page", "Small/medium tables that fit on one page"],
            ["`fit_width_only(ws)`", "Width fits 1 page; height grows; manual breaks honored", "Tables that are wide but span multiple pages tall"],
            ["`disable_fit_to_page(ws)`", "Turns off shrink-to-fit entirely", "When manual breaks should control everything"],
            ["`add_break_after(ws, row)`", "Inserts a horizontal page break after the given row", "Wherever you want a forced page boundary"],
        ],
        col_widths_inches=[2.0, 3.2, 1.8])
    add_callout(doc, "CRITICAL",
                "add_break_after only works if fit_single_page is OFF. If you call add_break_after, you almost always also call disable_fit_to_page (or fit_width_only). Otherwise Excel shrinks everything onto one page and the break has no effect.",
                color_hex="FFE8E8", border_hex="D88A8A")

    add_h2(doc, "6.3  The handler signature")
    add_code_block(doc, """\
def _handle_rule_NAME(ws, dest_filename):
    ...""")
    add_bullet(doc, "ws — the openpyxl Worksheet object for the sheet being processed.")
    add_bullet(doc, "dest_filename — absolute path of the .xlsx file. Useful for state-specific behavior.")
    add_para(doc, "The function returns nothing. It mutates ws in place.")

    add_h2(doc, "6.4  The openpyxl page-setup API — cheat sheet")
    add_table(doc,
        ["What you want", "Code"],
        [
            ["Set the print area", "`ws.print_area = f\"A1:H{ws.max_row}\"`"],
            ["Repeat top rows on every page", "`ws.print_title_rows = \"1:3\"`"],
            ["No print titles", "`ws.print_title_rows = None`"],
            ["Landscape orientation", "`ws.page_setup.orientation = \"landscape\"`"],
            ["Top margin (inches)", "`ws.page_margins.top = 1.00`"],
            ["Center horizontally", "`ws.print_options.horizontalCentered = True`"],
            ["Add page break after row N", "`add_break_after(ws, N)`"],
            ["Read a cell", "`ws[\"A10\"].value`"],
            ["Iterate rows", "`for row in range(1, ws.max_row + 1):`"],
        ],
        col_widths_inches=[2.6, 4.4])

    add_h2(doc, "6.5  Step-by-step: adding 'Rule 999'")
    add_para(doc, "Suppose Rule 999 sheets need print area A1:H{max_row}, manual page break every 50 rows, top margin 1 inch — but if cell B4 says \"Special\", just fit the sheet on one page instead.")

    add_h3(doc, "Step 1 — Write the handler in BApagebreaks.py")
    add_code_block(doc, """\
def _handle_rule_999(ws, dest_filename):
    # Special case: short sheet, just fit it
    if ws["B4"].value == "Special":
        fit_single_page(ws)
        return

    ws.print_area = f"A1:H{ws.max_row}"
    ws.page_margins.top = 1.00
    disable_fit_to_page(ws)
    for row in range(50, ws.max_row, 50):
        add_break_after(ws, row)""")

    add_h3(doc, "Step 2 — Register it in SHEET_RULES")
    add_code_block(doc, """\
SHEET_RULES = [
    ...
    ("Rule R1",   _handle_rule_r1),
    ("Rule 999",  _handle_rule_999),    # ← new line
]""")
    add_para(doc, "That is it. The next run, every sheet whose name starts with \"Rule 999\" gets this treatment.")

    add_h2(doc, "6.6  Recipes for common scenarios")

    add_h3(doc, "A) Whole sheet on one page")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    fit_single_page(ws)""")

    add_h3(doc, "B) Manual break after row 37")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    disable_fit_to_page(ws)
    add_break_after(ws, 37)""")

    add_h3(doc, "C) Break every 45 rows")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    disable_fit_to_page(ws)
    for row in range(45, ws.max_row, 45):
        add_break_after(ws, row)""")

    add_h3(doc, "D) Restrict print area to columns A–H")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    ws.print_area = f"A1:H{ws.max_row}"
    fit_single_page(ws)""")

    add_h3(doc, "E) Repeat rows 1–3 at the top of every printed page")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    fit_width_only(ws)
    ws.print_title_rows = "1:3" """)

    add_h3(doc, "F) Landscape")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    ws.page_setup.orientation = "landscape"
    fit_width_only(ws)""")

    add_h3(doc, "G) Break BEFORE every row whose column-A text starts with 'Section '")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    disable_fit_to_page(ws)
    for row in range(1, ws.max_row + 1):
        value = str(ws.cell(row=row, column=1).value or "")
        if value.startswith("Section ") and row > 3:
            ws.row_breaks.append(Break(id=row - 1))   # break BEFORE row""")

    add_h3(doc, "H) State-specific tweak (only for FL)")
    add_code_block(doc, """\
def _handle_my_rule(ws, dest_filename):
    fit_single_page(ws)
    if "FL" not in dest_filename:
        ws.page_margins.top = 1.00""")

    add_h2(doc, "6.7  Rule precedence — order matters")
    add_para(doc, "SHEET_RULES is checked in order. The FIRST matching prefix wins. Always list more-specific prefixes BEFORE less-specific ones:")
    add_code_block(doc, """\
SHEET_RULES = [
    ("Rule 239 C",   _handle_rule_239c),         # specific  ← FIRST
    ("Rule 239 ",    _handle_rule_239_general),  # generic   ← AFTER
]""")
    add_para(doc, "If you swap the order, \"Rule 239 C\" sheets would match the generic handler and never reach the specific one.")
    add_para(doc, "A sheet that matches no rule keeps the defaults (single-page fit, row 1 repeats).")

    add_h2(doc, "6.8  Bonus: the XML sanitize pass")
    add_para(doc, "After saving, _sanitize_xlsx opens the .xlsx zip and removes a few quirks openpyxl leaves behind (invalid definedName entries with $0 refs). Without this, Excel pops up \"We found a problem with some content. Open and Repair?\" on every file open. You do not need to do anything — it runs automatically.")

    # ====================================================================
    # SECTION 7 — PART 4: PDF GENERATION
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "7.  PART 4 — PDF generation")
    add_horizontal_rule(doc)

    add_h2(doc, "7.1  How it works")
    add_bullet(doc, "The .xlsx already has all print settings baked in (page breaks, fit-to-page, print areas, headers, footers — all set during the page-breaks pass).")
    add_bullet(doc, "export_to_pdf launches Excel via COM (win32com), opens the workbook read-only, hides the Index tab so it does not appear in the PDF, then calls Excel's own ExportAsFixedFormat to produce the PDF.")
    add_bullet(doc, "The function then verifies the PDF actually exists on disk and is non-empty. If not, it raises RuntimeError.")
    add_bullet(doc, "The PDF is saved next to the .xlsx with the same name. Example: \"ME 03-01-26 BA Small Market Rate Pages.pdf\".")

    add_h2(doc, "7.2  How the UI uses it")
    add_bullet(doc, "After Excel creation, the green \"Excel created\" message appears with a \"Generate PDF Document\" button.")
    add_bullet(doc, "Clicking the button calls generate_pdf_only. On success, a second green message appears: \"PDF created: <filename>\".")
    add_bullet(doc, "On failure, the error is shown in red and the button stays available so you can retry.")

    add_h2(doc, "7.3  Requirements")
    add_bullet(doc, "Microsoft Excel must be installed on the machine running the app — the COM bridge talks to a real Excel process.")
    add_bullet(doc, "pywin32 must be installed in the venv (already installed).")
    add_bullet(doc, "Close the file in Excel before clicking \"Generate PDF\" — if it is open elsewhere, Excel will refuse to overwrite.")

    add_h2(doc, "7.4  Multi-state mode")
    add_para(doc, "In multi-state mode, ticking \"Generate PDF for each state\" runs generate_pdf_only once per state, saving each PDF into a PDF/ subfolder inside the save location.")

    add_h2(doc, "7.5  Customizing the PDF")
    add_para(doc, "The PDF mirrors whatever the .xlsx print settings say. Customizations:")
    add_bullet(doc, "Want an extra page break? Add it in the page-breaks rule (Section 6).")
    add_bullet(doc, "Want different margins? Set ws.page_margins.* in the page-breaks rule.")
    add_bullet(doc, "Want landscape? ws.page_setup.orientation = \"landscape\" in the page-breaks rule.")
    add_bullet(doc, "Want the Index page to appear in the PDF? Remove the Index-hide block in export_to_pdf.")
    add_callout(doc, "PRINCIPLE",
                "You never touch PDF-specific code to customize layout. Every layout decision lives in the page-breaks layer. The PDF export just respects what is already there.",
                color_hex="EAF1FB", border_hex="9FB7DA")

    # ====================================================================
    # SECTION 8 — COMMON MISTAKES
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "8.  Common mistakes (read this!)")
    add_horizontal_rule(doc)

    mistakes = [
        ("Mistake 1 — Manual breaks while fit-to-page is on",
         "fit_single_page shrinks everything to one page; manual breaks become no-ops.",
         """# WRONG
def _handle_my_rule(ws, dest_filename):
    fit_single_page(ws)
    add_break_after(ws, 37)    # ineffective!

# RIGHT
def _handle_my_rule(ws, dest_filename):
    disable_fit_to_page(ws)
    add_break_after(ws, 37)"""),

        ("Mistake 2 — Listing a generic rule prefix before a specific one",
         "The first matching prefix wins. \"Rule 222\" eats \"Rule 222 B\" and \"Rule 222 TTT\" if listed first.",
         """# WRONG order
SHEET_RULES = [
    ("Rule 222",     _handle_rule_222_generic),
    ("Rule 222 TTT", _handle_rule_222ttt),    # never runs!
]

# RIGHT: specific first
SHEET_RULES = [
    ("Rule 222 TTT", _handle_rule_222ttt),
    ("Rule 222 B",   _handle_rule_222b),
    ("Rule 222",     _handle_rule_222_generic),
]"""),

        ("Mistake 3 — Forgetting overideFooter",
         "Without it, a clustered tab keeps the default footer (which assumes all 4 companies). The DoI filing will list the wrong names.",
         None),

        ("Mistake 4 — Wrong useIndex value",
         "useIndex=True writes pandas' 0,1,2,3,... as the first column. Almost always you want False.",
         """# WRONG
RatePages.generateWorksheet('Rule 999', 'TITLE', 'subtitle', df, True, True)

# RIGHT
RatePages.generateWorksheet('Rule 999', 'TITLE', 'subtitle', df, False, True)"""),

        ("Mistake 5 — add_break_after off-by-one",
         "add_break_after(ws, N) adds a break AFTER row N (the next page starts on row N+1). Use the helper instead of constructing Break(id=...) manually.",
         None),

        ("Mistake 6 — Mutating the workbook after process_pagebreaks runs",
         "Pipeline order is: build → save → page-break → save → PDF. Changes made after step 3 will not be in the saved file unless you save again.",
         None),

        ("Mistake 7 — Trying to generate a PDF without Excel installed",
         "export_to_pdf requires Excel + pywin32. There is no fallback — Excel's PDF engine is the only thing that respects every print setting accurately.",
         None),

        ("Mistake 8 — Editing _apply_default_footer for a one-off change",
         "That method runs for every sheet. For a single-rule override, modify the worksheet directly after generateWorksheet*:",
         """RatePages.generateWorksheet('Rule 999', 'TITLE', 'subtitle', df, False, True)
ws = RatePages.getWB()['Rule 999']
ws.oddFooter.left.text = "Custom footer text"   # one-off override"""),
    ]

    for title, desc, code in mistakes:
        add_h3(doc, title)
        add_para(doc, desc)
        if code:
            add_code_block(doc, code)

    # ====================================================================
    # SECTION 9 — CHEAT SHEET
    # ====================================================================
    doc.add_page_break()
    add_h1(doc, "9.  Cheat sheet — where to find what")
    add_horizontal_rule(doc)
    add_table(doc,
        ["Concern", "File / Location"],
        [
            ["Add a new rule (new tab)", "BARates.py → buildBAPages()"],
            ["Change what data goes on a sheet", "BARates.py → the build* method for that rule"],
            ["Change tab title / subtitle", "BARates.py → arguments to generateWorksheet*"],
            ["Change layout shape (1 table → 3 tables)", "BARates.py → switch to a different generateWorksheet* method"],
            ["Add a new layout shape", "ExcelSettingsBA.py → mimic generate_stacked_tables"],
            ["Change cell styling (borders, fonts, widths)", "ExcelSettingsBA.py → matching format* method"],
            ["Change top header text", "ExcelSettingsBA.py → _apply_standard_header"],
            ["Change left header text globally", "config/constants.py → HEADER_LEFT_TEXT"],
            ["Change default footer logic", "ExcelSettingsBA.py → _apply_default_footer"],
            ["Override footer for one cluster", "BARates.py → self.overideFooter(ws, CompanyTest)"],
            ["Add page-break rule", "BApagebreaks.py → write handler + add to SHEET_RULES"],
            ["Change page-break defaults (all sheets)", "BApagebreaks.py → top of for-loop in process_pagebreaks"],
            ["PDF generation", "BApagebreaks.py → export_to_pdf; called by BARatePages.generate_pdf_only"],
            ["Fonts, margins, paths, company names", "config/constants.py"],
            ["State-specific sheet selection", "BA Input File.xlsx (read by sheet_fetch in buildBAPages)"],
            ["Streamlit UI behavior", "app.py"],
        ],
        col_widths_inches=[3.4, 3.6])

    add_h1(doc, "TL;DR")
    add_horizontal_rule(doc)
    add_h3(doc, "Adding a new rule")
    add_code_block(doc, """\
# In BARates.py → buildBAPages()
self.compareCompanies('YourTable_Ext')
for CompanyTest in self.CompanyListDif:
    comp_name = self.extract_company_name(CompanyTest)
    self.title_company_name = CompanyTest
    if len(self.CompanyListDif) == 1:
        self.title_company_name = ""

    RatePages.generateWorksheet(
        'Rule 999 ' + self.title_company_name,
        'RULE 999. YOUR TITLE ' + self.title_company_name,
        '999.A. Your Subtitle',
        self.buildRule999(comp_name),
        False, True
    )
    self.overideFooter(RatePages.getWB()['Rule 999 ' + self.title_company_name], CompanyTest)

# At the bottom of buildBAPages, add a format pass:
Rule999_Sheets = [n for n in excel_Sheet_names if n.startswith('Rule 999')]
for r in Rule999_Sheets:
    self.formatWorksheet(AutoPages[r])""")

    add_h3(doc, "Adding a page-break rule")
    add_code_block(doc, """\
# In BApagebreaks.py
def _handle_rule_999(ws, dest_filename):
    disable_fit_to_page(ws)
    add_break_after(ws, 37)

SHEET_RULES = [
    ...
    ("Rule 999", _handle_rule_999),
]""")

    add_h3(doc, "Generating the PDF")
    add_para(doc, "Click \"Generate PDF Document\" in the UI. The export_to_pdf function uses Excel's own engine, preserves every print setting, saves next to the xlsx, and raises if anything fails.")

    # ---- Footer note ----
    add_para(doc, "")
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nationwide Insurance  ·  BA Analytics Division  ·  Internal Use Only")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = TEXT_MUTED

    out_path = "BA_Rate_Pages_Developer_Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
